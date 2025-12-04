import os
from flask import Flask, request, jsonify, send_file
from google.genai import Client
import json
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

app = Flask(__name__)

# --- CONFIGURACIÓN DE GEMINI API (¡VERSIÓN SEGURA!) ---
API_KEY = os.environ.get("GEMINI_API_KEY")

# --- CONFIGURACIÓN DE GEMINI API (VERSIÓN SEGURA) ---
API_KEY = os.environ.get("GEMINI_API_KEY")

if not API_KEY:
    raise ValueError(
        "ERROR: No se encontró la variable de entorno GEMINI_API_KEY.\n"
        "En Render debes configurarla en: Settings → Environment → Add Environment Variable.\n"
        "Clave: GEMINI_API_KEY | Valor: tu_api_key"
    )

try:
    client = Client(api_key=API_KEY)
except Exception as e:
    raise RuntimeError(f"ERROR: No se pudo inicializar el cliente Gemini. Detalle: {e}")


try:
    client = Client(api_key=API_KEY)
except Exception as e:
    print(f"ERROR: No se pudo crear el cliente Gemini. Detalle: {e}")

# Variable global para almacenar la última sesión generada
ultima_sesion = None

# ...existing code...
@app.route('/')
def index():
    try:
        with open('index.html', 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return "Error: No se encuentra el archivo index.html.", 500
# ...existing code...


@app.route('/suggest', methods=['POST'])
def generar_sugerencias():
    """Genera una lista de 3-5 sugerencias para un campo específico usando Gemini."""
    try:
        datos = request.json
        campo_solicitado = datos.get('campo') # 'competencia', 'capacidad', 'desempeno', etc.
        tema = datos.get('tema', 'N/A')
        nivel = datos.get('nivel', 'N/A')
        grado = datos.get('grado', 'N/A')
        area = datos.get('area', 'N/A')

        if not campo_solicitado or tema == 'N/A':
            return jsonify({'error': 'Faltan datos de entrada para la sugerencia (Tema, Nivel, Grado, Área).'}), 400

    
        prompt = f"""
Eres un experto en diseño curricular del MINEDU (Perú). Basado en el Tema: '{tema}', Nivel: '{nivel}', Grado: '{grado}', y Área: '{area}', genera una lista de 3 a 5 opciones de {campo_solicitado} relevantes al currículo nacional.

IMPORTANTE: Responde ÚNICAMENTE con un objeto JSON válido, sin texto adicional antes o después. El JSON debe contener una clave llamada '{campo_solicitado}_sugerencias' con una lista de strings.

Ejemplo de JSON esperado:
{{
    "{campo_solicitado}_sugerencias": [
        "Opción 1 de {campo_solicitado}.",
        "Opción 2 de {campo_solicitado}.",
        "Opción 3 de {campo_solicitado}."
    ]
}}
"""
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )

        response_text = response.text.strip()
        
        # Limpieza de la respuesta para el JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:].strip()
        if response_text.endswith('```'):
            response_text = response_text[:-3].strip()

        sugerencias_data = json.loads(response_text)
        
        # Verifica que la IA haya devuelto la clave esperada
        key = campo_solicitado + '_sugerencias'
        if key not in sugerencias_data:
             raise ValueError("La IA no devolvió la lista esperada en el formato correcto.")

        return jsonify({'sugerencias': sugerencias_data[key]})

    except json.JSONDecodeError:
        return jsonify({'error': 'Error de formato JSON en la respuesta de IA. Intente de nuevo.'}), 500
    except Exception as e:
        return jsonify({'error': f'Error al generar sugerencias (API/Servidor): {str(e)}'}), 500


@app.route('/generate', methods=['POST'])
def generar_sesion():
    global ultima_sesion

    try:
        datos = request.json

        # OBTENER TODOS LOS DATOS DEL FORMULARIO
        tema = datos.get('tema', 'Tema no especificado')
        nivel = datos.get('nivel', 'N/A')
        grado = datos.get('grado', 'N/A')
        area = datos.get('area', 'N/A')
        competencia = datos.get('competencia', 'N/A')
        capacidad = datos.get('capacidad', 'N/A')
        desempeno = datos.get('desempeno', 'N/A')
        comp_transversal = datos.get('comp_transversal', 'N/A')
        cap_transversal = datos.get('cap_transversal', 'N/A')
        enfoque = datos.get('enfoque', 'N/A')
        valor = datos.get('valor', 'N/A')
        
        try:
            tiempo = int(datos.get('tiempo', 90))
        except ValueError:
            tiempo = 90
            
        inicio_min = round(tiempo * 0.2)
        desarrollo_min = round(tiempo * 0.6)
        cierre_min = round(tiempo * 0.2)


        # --- PROMPT MAESTRO OPTIMIZADO PARA CONTENIDO CONCISO ---
        prompt = f"""
Eres un experto en diseño curricular del MINEDU (Perú). Genera una Sesión de Aprendizaje basada en los siguientes datos.
IMPORTANTE: Debes responder ÚNICAMENTE con un objeto JSON válido, sin texto adicional antes o después.

- Tema: {tema}
- Nivel: {nivel} y Grado: {grado}
- Área: {area}
- Tiempo total: {tiempo} minutos.
- Competencia principal: {competencia}
- Capacidad relacionada: {capacidad}
- Desempeño Principal: {desempeno}
- Competencia Transversal: {comp_transversal}
- Capacidad Transversal: {cap_transversal}
- Enfoque Transversal: {enfoque}
- Valor asociado: {valor}

El JSON debe tener exactamente esta estructura:
{{
    "titulo_sesion": "Genera un título atractivo y profesional para la sesión.",
    "proposito": "Descripción clara y concisa del propósito de aprendizaje (1-2 líneas).",
    "evidencia": "Descripción breve de la evidencia observable del aprendizaje.",
    "datos_adicionales": {{
        "competencia_transversal": "{comp_transversal}",
        "capacidad_transversal": "{cap_transversal}",
        "enfoque_transversal": "{enfoque}",
        "valor_asociado": "{valor}",
        "tiempo_total": "{tiempo} minutos"
    }},
    "criterios_evaluacion": [
        "Criterio 1 basado en el Desempeño y observable.",
        "Criterio 2 basado en la Competencia y medible."
    ],
    "secuencia_didactica": {{
        "inicio": "Descripción RESUMIDA y CLARA de las actividades de inicio ({inicio_min} min). Incluye: motivación, saberes previos y presentación del propósito. Máximo 80 palabras, usa viñetas o puntos clave.",
        "desarrollo": "Descripción RESUMIDA y CLARA de las actividades de desarrollo ({desarrollo_min} min). Incluye: pasos principales, uso de materiales y rol del docente. Máximo 120 palabras, usa viñetas o puntos clave.",
        "cierre": "Descripción RESUMIDA y CLARA de las actividades de cierre ({cierre_min} min). Incluye: metacognición, reflexión y evaluación. Máximo 60 palabras, usa viñetas o puntos clave."
    }}
}}
"""

        # Llamada a la API de Gemini
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )

        response_text = response.text.strip()

        # Limpieza de la respuesta para el JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:].strip()
        if response_text.endswith('```'):
            response_text = response_text[:-3].strip()

        sesion_data = json.loads(response_text)

        # Guardar la sesión para la descarga
        ultima_sesion = {
            'datos_form': datos,
            'sesion': sesion_data
        }

        # Devolver el JSON de la sesión
        return jsonify({'sesion': sesion_data})

    except json.JSONDecodeError as e:
        error_msg = f"Error de formato JSON en la respuesta de IA: {str(e)}"
        print(f"TRACEBACK: {error_msg}")
        return jsonify({'error': 'La IA devolvió un formato incorrecto. Revise la terminal.'}), 500

    except Exception as e:
        error_msg = f"Error al generar la sesión (API/Servidor): {str(e)}"
        print(f"TRACEBACK: {error_msg}")
        return jsonify({'error': error_msg}), 500

@app.route('/download')
def descargar_word():
    global ultima_sesion

    if not ultima_sesion:
        return jsonify({'error': 'No hay sesión generada'}), 400

    try:
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        sesion = ultima_sesion['sesion']
        datos = ultima_sesion['datos_form']
        datos_adicionales = sesion.get('datos_adicionales', {})
        
        # --- Datos para Tiempos (Calculados) ---
        try:
            tiempo_total = int(datos_adicionales.get('tiempo_total', '90 minutos').split(' ')[0])
        except ValueError:
            tiempo_total = 90
            
        inicio_min_doc = round(tiempo_total * 0.2)
        desarrollo_min_doc = round(tiempo_total * 0.6)
        cierre_min_doc = round(tiempo_total * 0.2)
        
        # --- Función para aplicar sombreado a celdas ---
        def shade_cell(cell, shade_color):
            """Aplica color de fondo a una celda"""
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), shade_color)
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # --- Función Auxiliar para Texto con Negrita y Centrado ---
        def add_bold_centered_text(cell, text, font_size=10, shade_color=None):
            if shade_color:
                shade_cell(cell, shade_color)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
            
        # --- Función Auxiliar para Texto con Negrita (Alineado a la Izquierda) ---
        def add_bold_text(cell, text, font_size=10, shade_color=None):
            if shade_color:
                shade_cell(cell, shade_color)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
            
        # --- Función para texto normal ---
        def add_normal_text(cell, text, font_size=10, align='left'):
            p = cell.paragraphs[0]
            if align == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'justify':
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
            
        # --- Función Auxiliar para Lista de Criterios (como texto) ---
        def format_criterios(criterios_list):
            return "\n".join([f"• {c}" for c in criterios_list])
            
        # --- NUEVOS DATOS ADMINISTRATIVOS ---
        dre = datos.get('dre', 'San Martín')
        ugel = datos.get('ugel', 'San Martín')
        ie = datos.get('ie', 'N/A')
        distrito = datos.get('distrito', 'Tarapoto')
        seccion = datos.get('seccion', 'A, B, C, D')
        ciclo = datos.get('ciclo', 'IV')
        director = datos.get('director', 'N/A')
        docente = datos.get('docente', 'N/A')
        fecha = datos.get('fecha', 'N/A')
        duracion = datos.get('duracion', '90\'')


        # --- ENCABEZADO PRINCIPAL ---
        titulo_principal = doc.add_paragraph()
        titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = titulo_principal.add_run('SESIÓN DE APRENDIZAJE\nINNOVACIÓN PEDAGÓGICA')
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)
        run_titulo.font.name = 'Arial'
        
        doc.add_paragraph()
        
        # --- TÍTULO DE LA SESIÓN EN RECUADRO ---
        tabla_titulo = doc.add_table(rows=1, cols=1)
        tabla_titulo.style = 'Table Grid'
        cell_titulo = tabla_titulo.cell(0, 0)
        add_bold_centered_text(cell_titulo, sesion.get('titulo_sesion', 'El suelo y su composición'), 11)
        
        doc.add_paragraph()

        # --- DATOS INFORMATIVOS ---
        # Icono de datos informativos (simulado con texto)
        p_datos = doc.add_paragraph()
        run_icono = p_datos.add_run('☰ ')
        run_icono.bold = True
        run_icono.font.size = Pt(11)
        run_datos = p_datos.add_run('DATOS INFORMATIVOS:')
        run_datos.bold = True
        run_datos.font.size = Pt(11)
        run_datos.font.name = 'Arial'
        
        # Tabla de datos informativos (5 filas x 8 columnas)
        tabla_info = doc.add_table(rows=5, cols=8)
        tabla_info.style = 'Table Grid'
        
        # Fila 1: DRE, San Martín, UGEL, San Martín
        add_bold_text(tabla_info.cell(0, 0), 'DRE', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(0, 1), dre, 9)
        add_bold_text(tabla_info.cell(0, 2), 'UGEL', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(0, 3).merge(tabla_info.cell(0, 7)), ugel, 9)
        
        # Fila 2: Institución Educativa, Distrito
        add_bold_text(tabla_info.cell(1, 0), 'Institución Educativa', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(1, 1).merge(tabla_info.cell(1, 3)), ie, 9)
        add_bold_text(tabla_info.cell(1, 4), 'Distrito', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(1, 5).merge(tabla_info.cell(1, 7)), distrito, 9)
        
        # Fila 3: Área curricular, Grado, Sección, Duración
        add_bold_text(tabla_info.cell(2, 0), 'Área curricular', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(2, 1), datos.get('area', 'N/A'), 9)
        add_bold_text(tabla_info.cell(2, 2), 'Grado', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(2, 3), datos.get('grado', 'N/A'), 9)
        add_bold_text(tabla_info.cell(2, 4), 'Sección', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(2, 5), seccion, 9)
        add_bold_text(tabla_info.cell(2, 6), 'Duración', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(2, 7), duracion, 9)
        
        # Fila 4: Ciclo, Fecha, Director(a)
        add_bold_text(tabla_info.cell(3, 0), 'Ciclo', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(3, 1), ciclo, 9)
        add_bold_text(tabla_info.cell(3, 2), 'Fecha', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(3, 3).merge(tabla_info.cell(3, 5)), fecha, 9)
        add_bold_text(tabla_info.cell(3, 6), 'Director(a)', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(3, 7), director, 9)
        
        # Fila 5: Docente
        add_bold_text(tabla_info.cell(4, 0), 'Docente', 9, 'D9D9D9')
        add_normal_text(tabla_info.cell(4, 1).merge(tabla_info.cell(4, 7)), docente, 9)
        
        doc.add_paragraph()

        # --- PROPÓSITO DE APRENDIZAJE ---
        p_proposito = doc.add_paragraph()
        run_icono2 = p_proposito.add_run('☰ ')
        run_icono2.bold = True
        run_icono2.font.size = Pt(11)
        run_proposito = p_proposito.add_run('PROPÓSITO DE APRENDIZAJE')
        run_proposito.bold = True
        run_proposito.font.size = Pt(11)
        run_proposito.font.name = 'Arial'
        
        tabla_proposito = doc.add_table(rows=2, cols=5)
        tabla_proposito.style = 'Table Grid'
        
        # Encabezados
        add_bold_centered_text(tabla_proposito.cell(0, 0), 'COMPETENCIA', 9, 'D9D9D9')
        add_bold_centered_text(tabla_proposito.cell(0, 1), 'CAPACIDAD', 9, 'D9D9D9')
        add_bold_centered_text(tabla_proposito.cell(0, 2), 'DESEMPEÑOS\nPRECISADOS', 9, 'D9D9D9')
        add_bold_centered_text(tabla_proposito.cell(0, 3), 'CRITERIOS DE EVALUACIÓN', 9, 'D9D9D9')
        add_bold_centered_text(tabla_proposito.cell(0, 4), 'INSTRUMENTO\nDE EVALUACIÓN', 9, 'D9D9D9')
        
        # Contenido
        row_data = tabla_proposito.rows[1].cells
        add_normal_text(row_data[0], datos.get('competencia', 'N/A'), 9, 'justify')
        add_normal_text(row_data[1], datos.get('capacidad', 'N/A'), 9, 'justify')
        add_normal_text(row_data[2], datos.get('desempeno', 'N/A'), 9, 'justify')
        
        criterios_formateados = format_criterios(sesion.get('criterios_evaluacion', ['N/A']))
        add_normal_text(row_data[3], criterios_formateados, 9, 'justify')
        add_normal_text(row_data[4], 'Lista de cotejo', 9, 'center')
        
        # Competencia transversal si existe
        comp_transversal_val = datos.get('comp_transversal', 'N/A')
        if comp_transversal_val not in ['N/A', 'No seleccionado', '']:
            row_transversal = tabla_proposito.add_row().cells
            add_normal_text(row_transversal[0], comp_transversal_val, 9, 'justify')
            add_normal_text(row_transversal[1], datos.get('cap_transversal', 'N/A'), 9, 'justify')
            add_normal_text(row_transversal[2], 'Integrada al desarrollo', 9, 'justify')
            add_normal_text(row_transversal[3], 'Observación de la capacidad', 9, 'justify')
            add_normal_text(row_transversal[4], 'Observación', 9, 'center')

        doc.add_paragraph()

        # --- ENFOQUES TRANSVERSALES ---
        p_enfoque = doc.add_paragraph()
        run_icono3 = p_enfoque.add_run('☰ ')
        run_icono3.bold = True
        run_icono3.font.size = Pt(11)
        run_enfoque = p_enfoque.add_run('ENFOQUES TRANSVERSALES')
        run_enfoque.bold = True
        run_enfoque.font.size = Pt(11)
        run_enfoque.font.name = 'Arial'
        
        tabla_enfoques = doc.add_table(rows=2, cols=4)
        tabla_enfoques.style = 'Table Grid'
        
        # Encabezados
        add_bold_centered_text(tabla_enfoques.cell(0, 0), 'ENFOQUE\nTRANSVERSAL', 9, 'D9D9D9')
        add_bold_centered_text(tabla_enfoques.cell(0, 1), 'VALORES', 9, 'D9D9D9')
        add_bold_centered_text(tabla_enfoques.cell(0, 2), 'ACTITUDES', 9, 'D9D9D9')
        add_bold_centered_text(tabla_enfoques.cell(0, 3), 'ACCIONES OBSERVABLES', 9, 'D9D9D9')
        
        # Contenido
        add_normal_text(tabla_enfoques.cell(1, 0), datos.get('enfoque', 'N/A'), 9, 'justify')
        add_normal_text(tabla_enfoques.cell(1, 1), datos.get('valor', 'N/A'), 9, 'justify')
        add_normal_text(tabla_enfoques.cell(1, 2), 'Disposición para adaptarse a los cambios, modificando si fuera necesario la propia conducta para alcanzar determinados objetivos.', 9, 'justify')
        add_normal_text(tabla_enfoques.cell(1, 3), 'Docentes y estudiantes comparan y adquieren y emplean estrategias útiles para aumentar la eficacia de sus esfuerzos en el logro de objetivos que se proponen.', 9, 'justify')

        doc.add_paragraph()

        # --- SECUENCIA DIDÁCTICA ---
        p_secuencia = doc.add_paragraph()
        run_icono4 = p_secuencia.add_run('☰ ')
        run_icono4.bold = True
        run_icono4.font.size = Pt(11)
        run_secuencia = p_secuencia.add_run('SECUENCIA DIDÁCTICA')
        run_secuencia.bold = True
        run_secuencia.font.size = Pt(11)
        run_secuencia.font.name = 'Arial'
        
        # Tabla vertical: 4 columnas (MOMENTOS | ACTIVIDADES | MEDIOS | TIEMPO)
        tabla_secuencia = doc.add_table(rows=4, cols=4)
        tabla_secuencia.style = 'Table Grid'
        
        secuencia_didactica = sesion.get('secuencia_didactica', {})
        
        # FILA 1: Encabezados
        add_bold_centered_text(tabla_secuencia.cell(0, 0), 'MOMENTOS', 10, 'D9D9D9')
        add_bold_centered_text(tabla_secuencia.cell(0, 1), 'ACTIVIDADES DE APRENDIZAJE', 10, 'D9D9D9')
        add_bold_centered_text(tabla_secuencia.cell(0, 2), 'MEDIOS Y MATERIALES', 10, 'D9D9D9')
        add_bold_centered_text(tabla_secuencia.cell(0, 3), 'TIEMPO', 10, 'D9D9D9')
        
        # FILA 2: MOTIVACIÓN / INICIO
        add_bold_text(tabla_secuencia.cell(1, 0), 'MOTIVACIÓN', 9)
        add_normal_text(tabla_secuencia.cell(1, 1), secuencia_didactica.get('inicio', 'N/A'), 9, 'justify')
        add_normal_text(tabla_secuencia.cell(1, 2), 'Pizarra, plumones, mota, proyector, laptop, parlantes, video', 9, 'justify')
        add_bold_centered_text(tabla_secuencia.cell(1, 3), f"{inicio_min_doc}'", 10)
        
        # FILA 3: PROCESAMIENTO DE LA INFORMACIÓN / DESARROLLO
        add_bold_text(tabla_secuencia.cell(2, 0), 'PROCESAMIENTO DE LA\nINFORMACIÓN', 9)
        add_normal_text(tabla_secuencia.cell(2, 1), secuencia_didactica.get('desarrollo', 'N/A'), 9, 'justify')
        add_normal_text(tabla_secuencia.cell(2, 2), 'Material concreto, Fichas de trabajo, Recursos educativos (video/ppt/web)', 9, 'justify')
        add_bold_centered_text(tabla_secuencia.cell(2, 3), f"{desarrollo_min_doc}'", 10)
        
        # FILA 4: EVALUACIÓN / CIERRE
        add_bold_text(tabla_secuencia.cell(3, 0), 'EVALUACIÓN', 9)
        add_normal_text(tabla_secuencia.cell(3, 1), secuencia_didactica.get('cierre', 'N/A'), 9, 'justify')
        add_normal_text(tabla_secuencia.cell(3, 2), 'Cuaderno, preguntas de metacognición', 9, 'justify')
        add_bold_centered_text(tabla_secuencia.cell(3, 3), f"{cierre_min_doc}'", 10)

        doc.add_paragraph()
        
        # --- BIBLIOGRAFÍAS ---
        p_biblio = doc.add_paragraph()
        run_icono5 = p_biblio.add_run('☰ ')
        run_icono5.bold = True
        run_icono5.font.size = Pt(11)
        run_biblio = p_biblio.add_run('BIBLIOGRAFÍAS:')
        run_biblio.bold = True
        run_biblio.font.size = Pt(11)
        run_biblio.font.name = 'Arial'
        
        doc.add_paragraph('• Currículo Nacional de la Educación Básica (MINEDU)')
        doc.add_paragraph('• Programación Curricular de Educación Primaria/Secundaria')
        doc.add_paragraph('• Recursos Web usados en el Desarrollo de la Sesión')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # --- FIRMAS ---
        tabla_firmas = doc.add_table(rows=1, cols=2)
        tabla_firmas.rows[0].cells[0].width = Inches(3)
        tabla_firmas.rows[0].cells[1].width = Inches(3)
        
        p_director = tabla_firmas.cell(0, 0).paragraphs[0]
        p_director.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_director.add_run('_' * 30 + '\n')
        p_director.add_run('V. B. director')
        
        p_docente = tabla_firmas.cell(0, 1).paragraphs[0]
        p_docente.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_docente.add_run('_' * 30 + '\n')
        p_docente.add_run('Docente')

        # Guardar en memoria
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='Sesion_ALLIMA_PUNCHAW.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({'error': f'Error al crear el documento Word: {str(e)}'}), 500

@app.route('/styles.css')
def styles():
    return send_file('styles.css')

@app.route('/script.js')
def scripts():
    return send_file('script.js')

@app.route('/home.html')
def home():
    try:
        return send_file('home.html')
    except FileNotFoundError:
        return "Error: No se encuentra el archivo home.html.", 404

if __name__ == '__main__':
    PORT = int(os.environ.get('PORT', 5007))
    HOST = '127.0.0.1'
    app.run(debug=True, host=HOST, port=PORT)

